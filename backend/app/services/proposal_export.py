"""Proposal export service"""
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime

from app.models.proposal import Proposal
from app.models.opportunity import Opportunity


async def export_proposal_to_docx(
    db: AsyncSession,
    proposal_id: str,
    tenant_id: str,
) -> BytesIO:
    """Export proposal to Word document (.docx)"""
    # Get proposal with opportunity
    result = await db.execute(
        select(Proposal, Opportunity)
        .join(Opportunity, Proposal.opportunity_id == Opportunity.id)
        .where(
            Proposal.id == proposal_id,
            Proposal.tenant_id == tenant_id
        )
    )
    row = result.first()
    if not row:
        raise ValueError("Proposal not found")
    
    proposal, opportunity = row
    
    # Create Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading(proposal.name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Proposal metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Version: {proposal.version}").bold = True
    meta_para.add_run(f" | Phase: {proposal.current_phase.value.replace('_', ' ').title()}")
    meta_para.add_run(f" | Status: {proposal.status.title()}")
    meta_para.add_run(f" | Opportunity: {opportunity.name}")
    if opportunity.agency:
        meta_para.add_run(f" | Agency: {opportunity.agency}")
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    # Executive Summary
    if proposal.executive_summary:
        doc.add_heading('Executive Summary', 1)
        para = doc.add_paragraph(proposal.executive_summary)
        para_format = para.paragraph_format
        para_format.space_after = Pt(12)
        doc.add_paragraph()  # Spacing
    
    # Technical Approach
    if proposal.technical_approach:
        doc.add_heading('Technical Approach', 1)
        para = doc.add_paragraph(proposal.technical_approach)
        para_format = para.paragraph_format
        para_format.space_after = Pt(12)
        doc.add_paragraph()  # Spacing
    
    # Management Approach
    if proposal.management_approach:
        doc.add_heading('Management Approach', 1)
        para = doc.add_paragraph(proposal.management_approach)
        para_format = para.paragraph_format
        para_format.space_after = Pt(12)
        doc.add_paragraph()  # Spacing
    
    # Past Performance
    if proposal.past_performance:
        doc.add_heading('Past Performance', 1)
        para = doc.add_paragraph(proposal.past_performance)
        para_format = para.paragraph_format
        para_format.space_after = Pt(12)
        doc.add_paragraph()  # Spacing
    
    # Win Themes
    if proposal.win_themes:
        doc.add_heading('Win Themes', 1)
        for theme in proposal.win_themes:
            para = doc.add_paragraph(theme, style='List Bullet')
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)
        doc.add_paragraph()  # Spacing
    
    # Footer with timestamp
    footer_para = doc.add_paragraph()
    footer_para.add_run(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}").italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

