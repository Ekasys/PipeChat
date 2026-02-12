"""RFP endpoints for ekchat-api."""
from __future__ import annotations

from collections import Counter
import csv
import io
import json
import re
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, Query, status, UploadFile
from fastapi.responses import Response, StreamingResponse
from sqlalchemy import text

from app.ai import complete_chat, stream_chat_completion
from app.config import settings
from app.deps import RequestContext, require_feature
from app.schemas import (
    RfpCapabilityMatrixGenerateRequest,
    RfpSectionGenerateRequest,
    RfpShredDocumentGenerateRequest,
)
from app.storage import get_storage


router = APIRouter(prefix="/rfp")

_HISTORY_FILE_KIND = "rfp_history_file"
_ANALYZE_FILE_KIND = "rfp_analyze_file"
_STYLE_PROFILE_KIND = "rfp_style_profile"
_SECTIONS_SESSION_KIND = "rfp_sections_session"
_RESPONSE_KIND = "rfp_generated_response"

_RFP_SECTION_TEXT_MAX_CHARS = int(getattr(settings, "EKCHAT_RFP_SECTION_TEXT_MAX_CHARS", 80000) or 80000)
_RFP_SECTION_MAX_SECTIONS = int(getattr(settings, "EKCHAT_RFP_SECTION_MAX_SECTIONS", 120) or 120)


def _sse_event(event: str, payload: Dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=True)}\n\n"


def _serialize_datetime(value: Any) -> Optional[str]:
    if isinstance(value, datetime):
        return value.isoformat()
    return None


def _coerce_epoch_seconds(value: Any) -> Optional[int]:
    if value is None:
        return None
    if isinstance(value, datetime):
        return int(value.timestamp())
    if isinstance(value, (int, float)):
        return int(value)
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return None
        if raw.isdigit():
            try:
                return int(raw)
            except ValueError:
                return None
        try:
            dt = datetime.fromisoformat(raw.replace("Z", "+00:00"))
            return int(dt.timestamp())
        except ValueError:
            return None
    return None


def _slug(value: str, max_len: int = 48) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "-", value.strip().lower()).strip("-")
    cleaned = cleaned[:max_len]
    return cleaned or "user"


def _safe_filename(value: Optional[str], *, default: str) -> str:
    candidate = Path(value or "").name.strip()
    candidate = re.sub(r"[^A-Za-z0-9._-]+", "-", candidate).strip(".-")
    return candidate or default


def _chunk_text(value: str, max_chars: int = 260) -> List[str]:
    text_value = (value or "").strip()
    if not text_value:
        return []

    words = text_value.split()
    chunks: List[str] = []
    current = ""

    for word in words:
        candidate = word if not current else f"{current} {word}"
        if len(candidate) > max_chars and current:
            chunks.append(current)
            current = word
        else:
            current = candidate

    if current:
        chunks.append(current)

    return chunks


def _scope_kind(scope: str) -> str:
    if scope == "history":
        return _HISTORY_FILE_KIND
    if scope == "analyze":
        return _ANALYZE_FILE_KIND
    raise HTTPException(status_code=400, detail=f"Unknown scope '{scope}'")


def _scope_blob_prefix(ctx: RequestContext, scope: str) -> str:
    return f"tenant/{ctx.user.tenant_id}/user/{ctx.user.user_id}/rfp/{scope}/files"


def _decode_bytes(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _extract_text_from_bytes(filename: str, raw: bytes, *, max_chars: int = 60000) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            return text.strip()[:max_chars]
        except Exception:
            pass

    if suffix == ".docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(raw))
            text = "\n".join((paragraph.text or "") for paragraph in document.paragraphs)
            return text.strip()[:max_chars]
        except Exception:
            pass

    if suffix in {".csv", ".tsv"}:
        delimiter = "\t" if suffix == ".tsv" else ","
        text_value = _decode_bytes(raw)
        reader = csv.reader(io.StringIO(text_value), delimiter=delimiter)
        rows: List[str] = []
        for idx, row in enumerate(reader):
            if idx >= 400:
                break
            rows.append(" | ".join(str(cell or "") for cell in row))
        return "\n".join(rows)[:max_chars]

    if suffix == ".json":
        try:
            parsed = json.loads(_decode_bytes(raw))
            return json.dumps(parsed, indent=2)[:max_chars]
        except json.JSONDecodeError:
            return _decode_bytes(raw)[:max_chars]

    return _decode_bytes(raw).strip()[:max_chars]


def _extract_doc_metadata(name: str, text_value: str) -> Dict[str, Any]:
    return {
        "name": name,
        "word_count": len(text_value.split()),
        "char_count": len(text_value),
    }


def _extract_requirements(text_value: str, *, max_items: int = 24) -> List[str]:
    lines = [line.strip(" -\t") for line in text_value.splitlines()]
    candidates: List[str] = []
    keywords = (
        "shall",
        "must",
        "required",
        "requirement",
        "deliver",
        "provide",
        "compliance",
        "security",
        "performance",
    )

    for line in lines:
        clean = " ".join(line.split())
        if len(clean) < 25:
            continue
        lower = clean.lower()
        if any(keyword in lower for keyword in keywords):
            candidates.append(clean)

    if not candidates:
        sentences = re.split(r"(?<=[.!?])\s+", text_value)
        for sentence in sentences:
            clean = " ".join(sentence.split())
            if len(clean) >= 30:
                candidates.append(clean)

    deduped: List[str] = []
    seen = set()
    for item in candidates:
        key = item.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
        if len(deduped) >= max_items:
            break

    return deduped


def _count_words(text_value: str) -> int:
    if not text_value:
        return 0
    return len(re.findall(r"\b\w+\b", text_value))


def _resolve_section_target_words(primary_words: int, fallback_words: int, default_words: int = 250) -> int:
    words = primary_words if primary_words > 0 else fallback_words
    if words <= 0:
        words = default_words
    return max(80, min(1200, words))


def _format_section_continuity(
    sections: List[Dict[str, Any]],
    current_index: int,
    *,
    max_sections: int = 2,
    max_chars: int = 1800,
) -> str:
    if not sections or current_index <= 1:
        return ""

    prior_blocks: List[str] = []
    ordered = sorted(sections, key=lambda item: int(item.get("index") or 0), reverse=True)
    for section in ordered:
        idx = section.get("index")
        if not isinstance(idx, int) or idx >= current_index:
            continue
        draft = str(section.get("draft") or section.get("content") or "").strip()
        if not draft:
            continue
        title = section.get("title") or f"Section {idx}"
        prior_blocks.append(f"Section {idx}: {title}\n{draft}")
        if len(prior_blocks) >= max_sections:
            break

    if not prior_blocks:
        return ""

    combined = "\n\n".join(reversed(prior_blocks))
    if len(combined) > max_chars:
        combined = combined[:max_chars].rstrip()
    return combined


def _extract_requirement_ids(text_value: str, fallback_index: int) -> List[str]:
    matches: List[str] = []
    for match in re.finditer(r"\b(?:REQ|REQID|R)[-_ ]?(\d{1,4})\b", text_value or "", flags=re.I):
        try:
            num = int(match.group(1))
        except Exception:
            continue
        matches.append(f"REQ-{num:03d}")

    if matches:
        out: List[str] = []
        seen = set()
        for item in matches:
            if item in seen:
                continue
            seen.add(item)
            out.append(item)
        return out

    if (text_value or "").strip():
        return [f"REQ-{fallback_index:03d}"]
    return []


def _extract_rfp_excerpt(text_value: str, title: str, requirements: str, *, max_chars: int = 480) -> str:
    haystack = (text_value or "").strip()
    if not haystack:
        return ""
    lower = haystack.lower()
    for term in (title or "", requirements or ""):
        token = term.strip()
        if not token:
            continue
        idx = lower.find(token.lower())
        if idx != -1:
            start = max(0, idx - 160)
            end = min(len(haystack), idx + max_chars)
            return haystack[start:end].strip()
    return haystack[:max_chars].strip()


_SECTION_TOKEN_STOPWORDS = {
    "the",
    "a",
    "an",
    "and",
    "or",
    "to",
    "of",
    "in",
    "on",
    "for",
    "with",
    "is",
    "are",
    "was",
    "were",
    "be",
    "as",
    "at",
    "by",
    "from",
    "it",
    "this",
    "that",
    "these",
    "those",
    "shall",
    "must",
    "required",
    "requirement",
    "provide",
    "deliver",
}


def _section_tokens(text_value: str) -> List[str]:
    raw = re.findall(r"[a-z0-9]+", (text_value or "").lower())
    return [token for token in raw if len(token) > 1 and token not in _SECTION_TOKEN_STOPWORDS]


def _section_overlap_score(query_tokens: List[str], segment_text: str) -> float:
    if not query_tokens:
        return 0.0
    query = set(query_tokens)
    segment = set(_section_tokens(segment_text))
    if not segment:
        return 0.0
    return len(query.intersection(segment)) / max(1, len(query))


def _select_history_excerpt(
    history_text: str,
    title: str,
    requirements: str,
    *,
    max_chars: int = 2200,
    max_segments: int = 3,
) -> str:
    source = (history_text or "").strip()
    if not source:
        return ""

    segments = [seg.strip() for seg in re.split(r"\n{2,}", source) if seg.strip()]
    if not segments:
        return source[:max_chars].strip()

    query_tokens = _section_tokens(f"{title} {requirements}")
    scored: List[Tuple[float, str]] = []
    for segment in segments:
        score = _section_overlap_score(query_tokens, segment)
        if score > 0:
            scored.append((score, segment))

    if scored:
        scored.sort(key=lambda item: item[0], reverse=True)
        picked = [item[1] for item in scored[:max_segments]]
    else:
        picked = segments[:1]

    excerpt = "\n\n".join(picked).strip()
    if len(excerpt) > max_chars:
        excerpt = excerpt[:max_chars].rstrip()
    return excerpt


_RFP_SECTION_ID_RE_TEXT = (
    r"(?:[A-Za-z]{1,6}\.\d+(?:\.\d+)*(?:\([A-Za-z]\)|[A-Za-z])?|"
    r"[A-Za-z]{1,6}\d+(?:\.\d+)*(?:\([A-Za-z]\)|[A-Za-z])?|"
    r"\d+(?:\.\d+)*(?:\([A-Za-z]\)|[A-Za-z])?)"
)
_RFP_TASK_HEADING_RE = re.compile(rf"^\s*(?P<num>{_RFP_SECTION_ID_RE_TEXT})(?P<suffix>[.)]?)\s*(?P<rest>.*)$")
_RFP_HEADING_ID_RE = re.compile(rf"^\s*({_RFP_SECTION_ID_RE_TEXT})")
_RFP_SECTION_NUM_RE = re.compile(rf"^\s*({_RFP_SECTION_ID_RE_TEXT})")
_RFP_SECTION_SPLIT_RE = re.compile(
    r"(?=\b(?:[A-Za-z]{1,6}\.\d+(?:\.\d+){0,6}|[A-Za-z]{1,6}\d+(?:\.\d+){0,6}|\d+(?:\.\d+){1,6})[.)]?\s+[A-Z])"
)
_RFP_PWS_ANCHOR_RE = re.compile(
    r"\b(performance work statement|pws|performance requirements|statement of work|sow|scope of work)\b",
    flags=re.I,
)
_RFP_REQ_STARTS = (
    "the ",
    "contractor",
    "offeror",
    "vendor",
    "provider",
    "supplier",
    "government",
    "agency",
    "shall",
    "must",
    "will",
    "should",
)
_RFP_REQ_PHRASES = (
    " shall ",
    " must ",
    " will ",
    " should ",
    " is required ",
    " are required ",
    " responsible for ",
    " provide ",
    " deliver ",
    " perform ",
    " maintain ",
    " ensure ",
    " support ",
    " comply ",
    " coordinate ",
)
_RFP_REQ_ENTITY_PHRASES = (
    "the contractor",
    "the offeror",
    "the vendor",
    "the provider",
    "the supplier",
    "the government",
    "the agency",
)


def _heading_level_from_text(text_value: str) -> int:
    if not text_value:
        return 1
    match = _RFP_HEADING_ID_RE.match(text_value.strip())
    if not match:
        return 1
    parts = [part for part in match.group(1).split(".") if part]
    return max(1, len(parts))


def _normalize_section_prefix(num: str) -> str:
    raw = (num or "").strip()
    if not raw:
        return ""
    parts = [part for part in raw.split(".") if part]
    while parts and parts[-1] == "0":
        parts.pop()
    normalized = ".".join(parts) if parts else raw
    return normalized.lower()


def _section_number_from_title(title: str) -> str:
    match = _RFP_SECTION_NUM_RE.match(title or "")
    return match.group(1) if match else ""


def _section_matches_prefix(title: str, prefix: str) -> bool:
    if not prefix:
        return False
    number = _section_number_from_title(title)
    if not number:
        return False
    norm_prefix = _normalize_section_prefix(prefix)
    norm_number = _normalize_section_prefix(number)
    return norm_number == norm_prefix or norm_number.startswith(f"{norm_prefix}.")


def _looks_like_heading_title(text_value: str) -> bool:
    cleaned = (text_value or "").strip()
    if not cleaned:
        return False
    lower = cleaned.lower()
    if any(lower.startswith(prefix) for prefix in _RFP_REQ_ENTITY_PHRASES):
        return False
    words = [word for word in re.split(r"\s+", cleaned) if word]
    if not words or len(words) > 10:
        return False
    if any(word.lower().strip(",:;") in {"shall", "must", "will", "should"} for word in words):
        return False
    if cleaned.isupper():
        return True
    titled = sum(1 for word in words if word[:1].isupper())
    return titled / len(words) >= 0.6 or len(words) <= 4


def _looks_like_requirement(text_value: str) -> bool:
    cleaned = (text_value or "").strip()
    if not cleaned:
        return False
    lower = cleaned.lower()
    if any(lower.startswith(prefix) for prefix in _RFP_REQ_STARTS):
        return True
    padded = f" {lower} "
    if any(phrase in padded for phrase in _RFP_REQ_PHRASES):
        return True
    return len(lower.split()) >= 12 and lower.endswith((".", ";"))


def _split_title_prefix(rest: str) -> Optional[Tuple[str, str]]:
    if not rest:
        return None
    lower = rest.lower()
    for phrase in _RFP_REQ_ENTITY_PHRASES:
        idx = lower.find(phrase)
        if idx > 0:
            title = rest[:idx].strip()
            body = rest[idx:].strip()
            if _looks_like_heading_title(title):
                return title, body
    words = rest.split()
    for idx, word in enumerate(words):
        token = word.lower().strip(",:;")
        if token in {"shall", "must", "will", "should"} and idx > 0:
            title = " ".join(words[:idx])
            if _looks_like_heading_title(title):
                body = " ".join(words[idx:])
                return title, body
    return None


def _split_heading_line(raw_line: str, num: str, suffix: str, rest: str) -> Tuple[str, str]:
    if ":" in rest:
        left, right = rest.split(":", 1)
        if _looks_like_heading_title(left):
            heading = f"{num}{suffix} {left.strip()}".strip()
            return heading, right.strip()
    title_prefix = _split_title_prefix(rest)
    if title_prefix:
        title, body = title_prefix
        heading = f"{num}{suffix} {title}".strip()
        return heading, body.strip()
    if _looks_like_requirement(rest):
        return f"{num}{suffix}".strip(), rest.strip()
    return raw_line.strip(), ""


def _normalize_requirement_text(text_value: str) -> str:
    return (text_value or "").strip()


def _split_section_lines(text_value: str) -> str:
    if not text_value:
        return ""
    normalized = text_value.replace("\r", "\n")
    lines: List[str] = []
    for raw in normalized.split("\n"):
        line = raw.strip()
        if not line:
            lines.append("")
            continue
        parts = _RFP_SECTION_SPLIT_RE.split(line)
        for part in parts:
            chunk = part.strip()
            if chunk:
                lines.append(chunk)
    return "\n".join(lines).strip()


def _detect_pws_root_prefix(text_value: str) -> Optional[str]:
    for raw in (text_value or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        if not _RFP_PWS_ANCHOR_RE.search(line):
            continue
        match = _RFP_HEADING_ID_RE.match(line)
        if match:
            return _normalize_section_prefix(match.group(1))
        match = re.search(r"\bsection\s+([A-Za-z]{1,6}(?:\.\d+(?:\.\d+)*)?|\d+(?:\.\d+)*)", line, flags=re.I)
        if match:
            return _normalize_section_prefix(match.group(1))
        match = re.search(
            r"([A-Za-z]{1,6}(?:\.\d+(?:\.\d+)*)?|\d+(?:\.\d+)*)\s+.*(performance work statement|pws|performance requirements|statement of work|sow|scope of work)",
            line,
            flags=re.I,
        )
        if match:
            return _normalize_section_prefix(match.group(1))
    return None


def _extract_rfp_tasks(text_value: str, *, max_items: int = 0) -> List[Dict[str, Any]]:
    if not text_value:
        return []
    limit = max_items if max_items and max_items > 0 else 0
    tasks: List[Dict[str, Any]] = []
    current: Optional[Dict[str, Any]] = None

    for raw_line in text_value.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            if current is not None:
                current.setdefault("body_lines", []).append("")
            continue
        match = _RFP_TASK_HEADING_RE.match(stripped)
        if match:
            if current is not None:
                tasks.append(current)
                if limit and len(tasks) >= limit:
                    current = None
                    break
            num = match.group("num")
            suffix = match.group("suffix") or ""
            rest = (match.group("rest") or "").strip()
            heading, inline_text = _split_heading_line(stripped, num, suffix, rest)
            current = {
                "title": heading,
                "body_lines": [],
                "level": _heading_level_from_text(num),
            }
            if inline_text:
                current["body_lines"].append(inline_text)
            continue
        if current is not None:
            current.setdefault("body_lines", []).append(line)

    if current is not None and (not limit or len(tasks) < limit):
        tasks.append(current)

    cleaned: List[Dict[str, Any]] = []
    for item in tasks:
        heading = (item.get("title") or "").strip()
        body_lines = item.get("body_lines") or []
        while body_lines and not body_lines[0].strip():
            body_lines.pop(0)
        while body_lines and not body_lines[-1].strip():
            body_lines.pop()
        body = "\n".join(line.rstrip() for line in body_lines).rstrip()
        body = _normalize_requirement_text(body)
        if not heading:
            continue
        cleaned.append(
            {
                "title": heading,
                "requirements": body,
                "is_section": not bool(body),
                "level": item.get("level") or _heading_level_from_text(heading),
                "number": _section_number_from_title(heading),
            }
        )
        if limit and len(cleaned) >= limit:
            break
    return cleaned


def _build_sections_from_tasks(tasks: List[Dict[str, Any]], max_sections: int) -> List[Dict[str, str]]:
    cleaned: List[Dict[str, str]] = []
    seen = set()
    for item in tasks or []:
        title = (item.get("title") or "").strip()
        if not title:
            continue
        requirements = (item.get("requirements") or "").strip()
        key = f"{title.lower()}|{requirements.lower()}"
        if key in seen:
            continue
        seen.add(key)
        cleaned.append({"title": title, "requirements": requirements})
        if max_sections and len(cleaned) >= max_sections:
            break
    return cleaned


def _extract_pws_outline(text_value: str, *, max_sections: int) -> List[Dict[str, str]]:
    prepared = _split_section_lines(text_value or "")
    tasks = _extract_rfp_tasks(prepared, max_items=0)
    prefix = _detect_pws_root_prefix(prepared)
    if prefix:
        tasks = [task for task in tasks if _section_matches_prefix(str(task.get("title") or ""), prefix)]
    if tasks:
        outline = _build_sections_from_tasks(tasks, max_sections)
        if outline:
            return outline
    return []


def _heuristic_outline(text_value: str, *, max_sections: int = 10) -> List[Dict[str, str]]:
    candidates: List[str] = []
    seen = set()
    for raw in (text_value or "").splitlines():
        line = raw.strip()
        if not line or len(line) > 90:
            continue
        if re.match(r"^\d+(\.\d+)*\s+[A-Za-z].+", line):
            title = re.sub(r"^\d+(\.\d+)*\s+", "", line).strip()
        elif re.match(r"^[A-Z][A-Z0-9 &/\-]{3,}$", line):
            title = line.title()
        elif line.endswith(":") and len(line.split()) <= 8:
            title = line.rstrip(":").strip()
        elif len(line.split()) <= 6 and line[:1].isupper():
            title = line.strip()
        else:
            continue
        norm = re.sub(r"[^a-z0-9]+", "", title.lower())
        if not norm or norm in seen:
            continue
        seen.add(norm)
        candidates.append(title)
        if len(candidates) >= max_sections:
            break
    return [{"title": title, "requirements": ""} for title in candidates]


def _extract_rfp_response(text_value: str) -> str:
    if not text_value:
        return ""
    match = re.search(r"<response>(.*?)</response>", text_value, flags=re.S | re.I)
    if match:
        text_value = match.group(1)
    cleaned = text_value.strip()
    cleaned = re.sub(r"^\s*(here( is|'s)|sure|draft response|response:|rfp response:)\s*", "", cleaned, flags=re.I)
    skip_patterns = (
        r"^<<<.*>>>$",
        r"^(rfp|history excerpts|style profile)\s*:?\s*$",
    )
    lines: List[str] = []
    for line in cleaned.splitlines():
        stripped = line.strip()
        if not stripped:
            lines.append(line)
            continue
        if any(re.match(pattern, stripped, flags=re.I) for pattern in skip_patterns):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


async def _extract_rfp_outline(
    ctx: RequestContext,
    text_value: str,
    *,
    model_name: str,
    max_sections: int = 10,
) -> List[Dict[str, str]]:
    trimmed = (text_value or "")[:12000]
    if not trimmed:
        return []

    system_prompt = (
        "You extract a response outline from an RFP. Return JSON only. "
        "Output a JSON array of objects with keys: title, requirements. "
        "Use section titles that the vendor should respond to, derived from the RFP text. "
        "Keep requirements to 4-5 sentences."
    )
    user_prompt = f"RFP text:\n{trimmed}\n\nReturn JSON only."

    outline_text = ""
    try:
        outline_text = await complete_chat(
            ctx.db,
            ctx.user.tenant_id,
            [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            model_name=model_name,
            max_tokens=max(700, settings.EKCHAT_CHAT_MAX_TOKENS),
            temperature=0.0,
        )
    except Exception:
        outline_text = ""

    parsed: Any = None
    if outline_text:
        try:
            parsed = json.loads(outline_text)
        except Exception:
            match = re.search(r"\[(.*)\]", outline_text, flags=re.S)
            if match:
                try:
                    parsed = json.loads(f"[{match.group(1)}]")
                except Exception:
                    parsed = None

    if isinstance(parsed, dict):
        parsed = parsed.get("sections") or parsed.get("outline") or parsed.get("items") or []

    if isinstance(parsed, list):
        cleaned: List[Dict[str, str]] = []
        seen = set()
        for item in parsed:
            if isinstance(item, dict):
                title = str(item.get("title") or "").strip()
                requirements = str(item.get("requirements") or "").strip()
            else:
                title = str(item or "").strip()
                requirements = ""
            if not title:
                continue
            key = re.sub(r"[^a-z0-9]+", "", title.lower())
            if not key or key in seen:
                continue
            seen.add(key)
            cleaned.append({"title": title, "requirements": requirements})
            if len(cleaned) >= max_sections:
                break
        if cleaned:
            return cleaned

    return _heuristic_outline(text_value, max_sections=max_sections)


def _infer_capability_area(requirement: str) -> str:
    text_value = requirement.lower()
    if any(token in text_value for token in ("security", "cmmc", "nist", "access control")):
        return "Security"
    if any(token in text_value for token in ("cloud", "infrastructure", "hosting", "azure")):
        return "Infrastructure"
    if any(token in text_value for token in ("staff", "personnel", "team", "labor")):
        return "Staffing"
    if any(token in text_value for token in ("schedule", "timeline", "milestone", "deadline")):
        return "Program Management"
    if any(token in text_value for token in ("integration", "api", "interface", "data exchange")):
        return "Technical Integration"
    return "General"


def _match_tokens(text_value: str) -> List[str]:
    raw = re.findall(r"[a-z0-9]+", (text_value or "").lower())
    stop = {
        "the",
        "a",
        "an",
        "and",
        "or",
        "to",
        "of",
        "in",
        "on",
        "for",
        "with",
        "is",
        "are",
        "was",
        "were",
        "be",
        "as",
        "at",
        "by",
        "from",
        "it",
        "this",
        "that",
        "these",
        "those",
        "shall",
        "must",
        "required",
        "requirement",
        "provide",
        "deliver",
    }
    return [token for token in raw if token and token not in stop]


def _overlap_score(requirement_text: str, evidence_text: str) -> float:
    req_tokens = set(_match_tokens(requirement_text))
    if not req_tokens:
        return 0.0
    ev_tokens = set(_match_tokens(evidence_text))
    if not ev_tokens:
        return 0.0
    hits = len(req_tokens.intersection(ev_tokens))
    return hits / max(1, len(req_tokens))


def _short_snippet(text_value: str, *, max_chars: int = 240) -> str:
    cleaned = " ".join((text_value or "").split())
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[: max_chars - 1].rstrip() + "…"


def _coverage_from_score(score: float) -> int:
    if score >= 0.55:
        return 3
    if score >= 0.35:
        return 2
    if score > 0:
        return 1
    return 0


def _epoch_seconds(value: Any) -> Optional[int]:
    if isinstance(value, datetime):
        return int(value.timestamp())
    return None


async def _load_history_evidence_docs(
    ctx: RequestContext,
    *,
    max_files: int = 6,
    max_chars_per_file: int = 12000,
) -> List[Dict[str, str]]:
    docs: List[Dict[str, str]] = []
    history_files = await _list_scope_files(ctx, "history")
    if not history_files:
        return docs

    storage = get_storage()
    for item in history_files[:max_files]:
        blob_path = str(item.get("blob_path") or "").strip()
        name = str(item.get("name") or "history-document").strip()
        if not blob_path:
            continue
        try:
            raw = await storage.download_bytes(blob_path)
        except Exception:
            continue
        text_value = _extract_text_from_bytes(name, raw, max_chars=max_chars_per_file)
        if not text_value:
            continue
        docs.append({"name": name, "text": text_value})
    return docs


def _build_capability_rows(requirements: List[str], history_docs: List[Dict[str, str]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []

    for idx, requirement in enumerate(requirements, start=1):
        matches: List[Dict[str, Any]] = []
        for doc in history_docs:
            raw_text = str(doc.get("text") or "")
            if not raw_text:
                continue
            segments = [seg.strip() for seg in re.split(r"\n{2,}", raw_text) if seg.strip()]
            if not segments:
                segments = [raw_text]
            for segment in segments[:120]:
                score = _overlap_score(requirement, segment)
                if score <= 0:
                    continue
                matches.append(
                    {
                        "doc_id": str(doc.get("name") or "Reference"),
                        "snippet": _short_snippet(segment, max_chars=260),
                        "similarity": score,
                    }
                )

        matches.sort(key=lambda item: float(item.get("similarity") or 0.0), reverse=True)
        top = matches[:3]
        best_score = float(top[0]["similarity"]) if top else 0.0
        coverage = _coverage_from_score(best_score)

        evidence_excerpts = "\n".join(
            f"{item['doc_id']}: {item['snippet']}"
            for item in top
        ) or "No directly matching historical evidence was found."
        evidence_sources = "\n".join(
            f"{item['doc_id']} / p.N/A"
            for item in top
        ) or "None"

        if coverage >= 3:
            rationale = "Historical content strongly aligns with this requirement."
            gaps_actions = "Tighten wording to mirror requirement language and include measurable outcomes."
        elif coverage == 2:
            rationale = "Historical evidence partially aligns, but traceability to all obligations is incomplete."
            gaps_actions = "Add explicit compliance statements and map evidence to each obligation sentence."
        elif coverage == 1:
            rationale = "Only limited evidence overlap was found in historical submissions."
            gaps_actions = "Add targeted narrative, staffing, and technical proof points specific to this clause."
        else:
            rationale = "No supporting evidence was found in the uploaded history library."
            gaps_actions = "Draft net-new response text and supporting evidence for this requirement."

        rows.append(
            {
                "rfp_requirement_id": f"REQ-{idx:03d}",
                "capability_area": _infer_capability_area(requirement),
                "requirement_text": requirement,
                "clause_breakdown": requirement,
                "coverage_score": coverage,
                "rationale": rationale,
                "evidence_excerpts": evidence_excerpts,
                "evidence_sources": evidence_sources,
                "gaps_actions": gaps_actions,
                "references": [
                    {
                        "doc_id": item["doc_id"],
                        "source_page": None,
                        "evidence_snippet": item["snippet"],
                        "similarity": round(float(item.get("similarity") or 0.0), 4),
                    }
                    for item in top
                ],
            }
        )

    return rows


def _summary_keywords(text_value: str, *, max_items: int = 6) -> List[str]:
    tokens = _match_tokens(text_value)
    if not tokens:
        return []
    counts = Counter(tokens)
    ranked = [item for item, _ in counts.most_common(max_items * 2)]
    out: List[str] = []
    for token in ranked:
        if token in out:
            continue
        out.append(token)
        if len(out) >= max_items:
            break
    return out


def _build_shred_rows(text_value: str, *, max_rows: int = 80) -> List[Dict[str, Any]]:
    lines = [line.strip() for line in text_value.splitlines() if line.strip()]
    section = "General"
    rows: List[Dict[str, Any]] = []

    heading_pattern = re.compile(r"^(\d+(?:\.\d+)*\s+.+|[A-Z][A-Z0-9 /,&\-]{4,})$")
    keyword_pattern = re.compile(r"\b(shall|must|required|requirement|deliver|provide|compliance)\b", flags=re.I)

    row_idx = 0
    for line in lines:
        compact = " ".join(line.split())
        if heading_pattern.match(compact) and len(compact) <= 120:
            section = compact
            continue

        if len(compact) < 20:
            continue

        if keyword_pattern.search(compact):
            row_idx += 1
            row_id_match = re.match(r"^([A-Za-z]?\d+(?:\.\d+)*)\b", compact)
            row_id = row_id_match.group(1) if row_id_match else f"REQ-{row_idx:03d}"
            summary = _short_snippet(compact, max_chars=280)
            rows.append(
                {
                    "requirement_id": row_id,
                    "requirement_text": compact,
                    "summary": summary,
                    "keywords": _summary_keywords(compact),
                    "section": section,
                }
            )
            if len(rows) >= max_rows:
                break

    if not rows:
        for idx, requirement in enumerate(_extract_requirements(text_value, max_items=min(max_rows, 40)), start=1):
            rows.append(
                {
                    "requirement_id": f"REQ-{idx:03d}",
                    "requirement_text": requirement,
                    "summary": _short_snippet(requirement, max_chars=280),
                    "keywords": _summary_keywords(requirement),
                    "section": _infer_capability_area(requirement),
                }
            )

    return rows[:max_rows]


def _normalize_keywords(value: Any) -> List[str]:
    if isinstance(value, list):
        raw = value
    elif isinstance(value, str):
        try:
            parsed = json.loads(value)
            raw = parsed if isinstance(parsed, list) else value.split(",")
        except Exception:
            raw = re.split(r"[,;\n]+", value)
    else:
        raw = []

    out: List[str] = []
    seen = set()
    for item in raw:
        token = str(item or "").strip().strip(".,;:")
        if not token:
            continue
        key = token.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(token)
    return out


def _normalize_reference(ref: Any) -> Dict[str, Any]:
    if not isinstance(ref, dict):
        return {}
    doc_id = str(ref.get("doc_id") or ref.get("source") or "Reference").strip() or "Reference"
    snippet = str(ref.get("evidence_snippet") or ref.get("snippet") or "").strip()
    source_page = ref.get("source_page")
    if isinstance(source_page, str) and source_page.isdigit():
        source_page = int(source_page)
    elif not isinstance(source_page, int):
        source_page = None
    similarity = ref.get("similarity")
    if isinstance(similarity, str):
        try:
            similarity = float(similarity)
        except ValueError:
            similarity = None
    elif not isinstance(similarity, (int, float)):
        similarity = None
    return {
        "doc_id": doc_id,
        "source_page": source_page,
        "evidence_snippet": snippet,
        "similarity": round(float(similarity), 4) if isinstance(similarity, (int, float)) else None,
    }


def _normalize_capability_row(row: Dict[str, Any], idx: int) -> Dict[str, Any]:
    requirement_text = str(
        row.get("requirement_text")
        or row.get("clause_breakdown")
        or row.get("requirement")
        or ""
    ).strip()
    refs_raw = row.get("references")
    refs = []
    if isinstance(refs_raw, list):
        refs = [_normalize_reference(item) for item in refs_raw if isinstance(item, dict)]
        refs = [item for item in refs if item]
    coverage_value = row.get("coverage_score")
    if isinstance(coverage_value, str):
        try:
            coverage_value = int(float(coverage_value))
        except ValueError:
            coverage_value = None
    elif isinstance(coverage_value, float):
        coverage_value = int(round(coverage_value))
    elif not isinstance(coverage_value, int):
        coverage_value = None
    if coverage_value is not None:
        coverage_value = max(0, min(3, coverage_value))

    evidence_excerpts = str(
        row.get("evidence_excerpts")
        or row.get("evidence_summary")
        or ""
    ).strip()
    evidence_sources = str(
        row.get("evidence_sources")
        or row.get("evidence_source")
        or ""
    ).strip()
    if refs and not evidence_excerpts:
        evidence_excerpts = "\n".join(
            f"{ref.get('doc_id')}: {_short_snippet(str(ref.get('evidence_snippet') or ''), max_chars=260)}"
            for ref in refs
            if ref.get("evidence_snippet")
        ).strip()
    if refs and not evidence_sources:
        evidence_sources = "\n".join(
            f"{ref.get('doc_id')} / p.{int(ref['source_page']) + 1}" if isinstance(ref.get("source_page"), int) else f"{ref.get('doc_id')} / p.N/A"
            for ref in refs
        ).strip()

    requirement_id = str(
        row.get("rfp_requirement_id")
        or row.get("requirement_id")
        or f"REQ-{idx:03d}"
    ).strip()
    capability_area = str(
        row.get("capability_area")
        or _infer_capability_area(requirement_text)
    ).strip() or "General"
    is_section = bool(row.get("is_section")) or not requirement_text
    if is_section:
        coverage_value = None
        evidence_excerpts = ""
        evidence_sources = ""
        refs = []

    return {
        "rfp_requirement_id": requirement_id,
        "capability_area": capability_area,
        "requirement_text": requirement_text,
        "clause_breakdown": str(row.get("clause_breakdown") or requirement_text).strip(),
        "coverage_score": coverage_value,
        "rationale": str(row.get("rationale") or row.get("clause_level_findings") or "").strip(),
        "evidence_excerpts": evidence_excerpts,
        "evidence_sources": evidence_sources,
        "gaps_actions": str(row.get("gaps_actions") or row.get("gaps") or "").strip(),
        "references": refs[:3],
        "is_section": is_section,
    }


def _normalize_shred_row(row: Dict[str, Any], idx: int) -> Dict[str, Any]:
    requirement_id = str(row.get("requirement_id") or row.get("section") or f"REQ-{idx:03d}").strip()
    requirement_text = str(
        row.get("requirement_text")
        or row.get("requirement")
        or row.get("summary")
        or ""
    ).strip()
    summary = str(row.get("summary") or row.get("requirement") or requirement_text).strip()
    return {
        "requirement_id": requirement_id or f"REQ-{idx:03d}",
        "requirement_text": requirement_text,
        "summary": summary,
        "keywords": _normalize_keywords(row.get("keywords")),
    }


def _normalize_capability_payload_rows(rows: Any) -> List[Dict[str, Any]]:
    if not isinstance(rows, list):
        return []
    return [_normalize_capability_row(item, idx) for idx, item in enumerate(rows, start=1) if isinstance(item, dict)]


def _normalize_shred_payload_rows(rows: Any) -> List[Dict[str, Any]]:
    if not isinstance(rows, list):
        return []
    return [_normalize_shred_row(item, idx) for idx, item in enumerate(rows, start=1) if isinstance(item, dict)]


def _csv_export(rows: List[Dict[str, Any]], headers: List[Tuple[str, str]]) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow([header for header, _ in headers])

    for row in rows:
        writer.writerow([row.get(key, "") for _, key in headers])

    return output.getvalue().encode("utf-8")


def _xlsx_export(rows: List[Dict[str, Any]], headers: List[Tuple[str, str]]) -> bytes:
    try:
        from openpyxl import Workbook
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Excel export is unavailable: {exc}") from exc

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Export"
    sheet.append([header for header, _ in headers])

    for row in rows:
        values: List[Any] = []
        for _, key in headers:
            value = row.get(key, "")
            if isinstance(value, list):
                value = ", ".join(str(item) for item in value)
            elif isinstance(value, dict):
                value = json.dumps(value, ensure_ascii=False)
            values.append(value)
        sheet.append(values)

    stream = io.BytesIO()
    workbook.save(stream)
    return stream.getvalue()


async def _insert_session(
    ctx: RequestContext,
    *,
    kind: str,
    payload: Dict[str, Any],
    record_id: Optional[str] = None,
    status_value: str = "ready",
    expires_in_hours: Optional[int] = None,
) -> str:
    sid = record_id or str(uuid.uuid4())
    expires_at = datetime.utcnow() + timedelta(hours=expires_in_hours) if expires_in_hours else None

    await ctx.db.execute(
        text(
            """
            INSERT INTO ekchat.rfp_sessions (id, tenant_id, user_id, kind, status, payload, created_at, expires_at)
            VALUES (
                :id,
                :tenant_id,
                :user_id,
                :kind,
                :status,
                CAST(:payload AS JSONB),
                NOW(),
                :expires_at
            )
            """
        ),
        {
            "id": sid,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "kind": kind,
            "status": status_value,
            "payload": json.dumps(payload),
            "expires_at": expires_at,
        },
    )

    return sid


async def _list_sessions(ctx: RequestContext, *, kind: str) -> List[Dict[str, Any]]:
    result = await ctx.db.execute(
        text(
            """
            SELECT id, status, payload, created_at, expires_at
            FROM ekchat.rfp_sessions
            WHERE tenant_id = :tenant_id
              AND user_id = :user_id
              AND kind = :kind
            ORDER BY created_at DESC
            """
        ),
        {
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "kind": kind,
        },
    )

    rows: List[Dict[str, Any]] = []
    for row in result.mappings().all():
        payload = row.get("payload")
        if isinstance(payload, str):
            try:
                payload = json.loads(payload)
            except json.JSONDecodeError:
                payload = {}
        elif payload is None:
            payload = {}

        rows.append(
            {
                "id": row["id"],
                "status": row.get("status"),
                "payload": payload,
                "created_at": row.get("created_at"),
                "expires_at": row.get("expires_at"),
            }
        )

    return rows


async def _latest_session(ctx: RequestContext, *, kind: str) -> Optional[Dict[str, Any]]:
    sessions = await _list_sessions(ctx, kind=kind)
    return sessions[0] if sessions else None


async def _delete_session(ctx: RequestContext, *, kind: str, session_id: str) -> None:
    await ctx.db.execute(
        text(
            """
            DELETE FROM ekchat.rfp_sessions
            WHERE id = :id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
              AND kind = :kind
            """
        ),
        {
            "id": session_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "kind": kind,
        },
    )


async def _save_style_profile(ctx: RequestContext, summary: str) -> Dict[str, Any]:
    await ctx.db.execute(
        text(
            """
            DELETE FROM ekchat.rfp_sessions
            WHERE tenant_id = :tenant_id
              AND user_id = :user_id
              AND kind = :kind
            """
        ),
        {
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "kind": _STYLE_PROFILE_KIND,
        },
    )

    payload = {
        "summary": summary,
        "updated_at": datetime.utcnow().isoformat(),
    }
    sid = await _insert_session(ctx, kind=_STYLE_PROFILE_KIND, payload=payload)
    payload["id"] = sid
    return payload


async def _load_style_profile(ctx: RequestContext) -> Optional[Dict[str, Any]]:
    row = await _latest_session(ctx, kind=_STYLE_PROFILE_KIND)
    if not row:
        return None

    payload = dict(row["payload"])
    payload["id"] = row["id"]
    return payload


def _rfp_chat_id(ctx: RequestContext) -> str:
    return f"rfp-history-{_slug(ctx.user.tenant_id, 24)}-{_slug(ctx.user.user_id, 24)}"


async def _ensure_history_chat(ctx: RequestContext, model: Optional[str], title: Optional[str]) -> Dict[str, Any]:
    chat_id = _rfp_chat_id(ctx)
    selected_model = (model or "").strip() or settings.EKCHAT_DEFAULT_MODEL
    selected_title = (title or "").strip() or "Chat with your documents"

    await ctx.db.execute(
        text(
            """
            INSERT INTO ekchat.chats (id, tenant_id, user_id, title, model, created_at, updated_at)
            VALUES (:id, :tenant_id, :user_id, :title, :model, NOW(), NOW())
            ON CONFLICT (id)
            DO UPDATE SET title = EXCLUDED.title, model = EXCLUDED.model, updated_at = NOW()
            """
        ),
        {
            "id": chat_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "title": selected_title,
            "model": selected_model,
        },
    )

    return {
        "id": chat_id,
        "title": selected_title,
        "model": selected_model,
    }


async def _list_scope_files(ctx: RequestContext, scope: str) -> List[Dict[str, Any]]:
    kind = _scope_kind(scope)
    sessions = await _list_sessions(ctx, kind=kind)

    files: List[Dict[str, Any]] = []
    for row in sessions:
        payload = row["payload"] if isinstance(row["payload"], dict) else {}
        files.append(
            {
                "id": row["id"],
                "name": payload.get("name", ""),
                "mime_type": payload.get("mime_type"),
                "size": payload.get("file_size"),
                "blob_path": payload.get("blob_path"),
                "created_at": _serialize_datetime(row.get("created_at")),
            }
        )

    return files


async def _find_scope_file(ctx: RequestContext, scope: str, *, name: str) -> Optional[Dict[str, Any]]:
    files = await _list_scope_files(ctx, scope)
    for item in files:
        if item.get("name") == name:
            return item
    return None


async def _history_text_excerpt(ctx: RequestContext, *, max_chars: int = 6000) -> str:
    files = await _list_scope_files(ctx, "history")
    if not files:
        return ""

    storage = get_storage()
    chunks: List[str] = []
    remaining = max_chars

    for item in files[:4]:
        blob_path = item.get("blob_path")
        if not blob_path:
            continue
        try:
            raw = await storage.download_bytes(blob_path)
        except Exception:
            continue

        text_value = _extract_text_from_bytes(item.get("name") or "document", raw, max_chars=min(remaining, 2500))
        if not text_value:
            continue

        chunks.append(f"Document: {item.get('name')}\n{text_value}")
        remaining -= len(text_value)
        if remaining <= 0:
            break

    return "\n\n".join(chunks)[:max_chars]


async def _get_sections_session(ctx: RequestContext, session_id: str) -> Dict[str, Any]:
    result = await ctx.db.execute(
        text(
            """
            SELECT id, status, payload, created_at
            FROM ekchat.rfp_sessions
            WHERE id = :id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
              AND kind = :kind
            LIMIT 1
            """
        ),
        {
            "id": session_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "kind": _SECTIONS_SESSION_KIND,
        },
    )

    row = result.mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Section session not found")

    payload = row.get("payload")
    if isinstance(payload, str):
        try:
            payload = json.loads(payload)
        except json.JSONDecodeError:
            payload = {}
    elif payload is None:
        payload = {}

    return {
        "id": row["id"],
        "status": row.get("status"),
        "payload": payload,
    }


async def _update_sections_payload(ctx: RequestContext, session_id: str, payload: Dict[str, Any]) -> None:
    await ctx.db.execute(
        text(
            """
            UPDATE ekchat.rfp_sessions
            SET payload = CAST(:payload AS JSONB), status = 'ready'
            WHERE id = :id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
              AND kind = :kind
            """
        ),
        {
            "id": session_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "kind": _SECTIONS_SESSION_KIND,
            "payload": json.dumps(payload),
        },
    )


@router.post("/history/chat")
async def rfp_history_chat(
    body: Optional[Dict[str, Any]] = Body(default=None),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    payload = body or {}
    chat = await _ensure_history_chat(
        ctx,
        model=payload.get("model"),
        title=payload.get("title"),
    )
    return {"chat": chat}


@router.post("/history/chat/clear")
async def rfp_history_chat_clear(
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    chat_id = _rfp_chat_id(ctx)

    delete_result = await ctx.db.execute(
        text(
            """
            DELETE FROM ekchat.messages
            WHERE chat_id = :chat_id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
            """
        ),
        {
            "chat_id": chat_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
        },
    )

    await ctx.db.execute(
        text(
            """
            UPDATE ekchat.chats
            SET updated_at = NOW()
            WHERE id = :chat_id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
            """
        ),
        {
            "chat_id": chat_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
        },
    )

    cleared = getattr(delete_result, "rowcount", 0) or 0
    return {"ok": True, "cleared": int(max(0, cleared))}


@router.post("/history/upload")
async def rfp_history_upload(
    files: List[UploadFile] = File(...),
    folder: Optional[str] = Form(default=None),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    _ = folder

    storage = get_storage()
    uploaded: List[Dict[str, Any]] = []

    for upload in files:
        data = await upload.read()
        if not data:
            continue

        safe_name = _safe_filename(upload.filename, default="document.bin")
        session_id = str(uuid.uuid4())
        blob_path = f"{_scope_blob_prefix(ctx, 'history')}/{session_id}-{safe_name}"

        await storage.upload_bytes(blob_path, data, content_type=upload.content_type)

        payload = {
            "name": safe_name,
            "mime_type": upload.content_type,
            "blob_path": blob_path,
            "file_size": len(data),
        }

        await _insert_session(
            ctx,
            kind=_HISTORY_FILE_KIND,
            payload=payload,
            record_id=session_id,
            status_value="ready",
        )

        uploaded.append({"id": session_id, "name": safe_name, "size": len(data)})

    return {"ok": True, "uploaded": uploaded}


@router.get("/history/list")
async def rfp_history_list(
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    return {"files": await _list_scope_files(ctx, "history")}


@router.get("/history/download")
async def rfp_history_download(
    name: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    file_row = await _find_scope_file(ctx, "history", name=name)
    if not file_row:
        raise HTTPException(status_code=404, detail="File not found")

    storage = get_storage()
    content = await storage.download_bytes(file_row["blob_path"])
    media_type = file_row.get("mime_type") or "application/octet-stream"

    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'inline; filename="{Path(name).name}"'},
    )


@router.get("/history/doc")
async def rfp_history_doc(
    name: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    file_row = await _find_scope_file(ctx, "history", name=name)
    if not file_row:
        raise HTTPException(status_code=404, detail="File not found")

    storage = get_storage()
    content = await storage.download_bytes(file_row["blob_path"])
    text_value = _extract_text_from_bytes(name, content)

    if not text_value:
        raise HTTPException(status_code=400, detail="Could not extract readable text from document")

    return {
        "name": name,
        "metadata": _extract_doc_metadata(name, text_value),
        "text": text_value,
        "truncated": len(text_value) >= 60000,
        "is_pdf": Path(name).suffix.lower() == ".pdf",
    }


@router.post("/history/delete")
async def rfp_history_delete(
    body: Dict[str, Any] = Body(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    name = str(body.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="File name required")

    file_row = await _find_scope_file(ctx, "history", name=name)
    if not file_row:
        raise HTTPException(status_code=404, detail="File not found")

    storage = get_storage()
    try:
        await storage.delete_path(file_row["blob_path"])
    except Exception:
        pass

    await _delete_session(ctx, kind=_HISTORY_FILE_KIND, session_id=file_row["id"])
    return {"ok": True, "files": await _list_scope_files(ctx, "history")}


@router.post("/history/style/edit")
async def rfp_history_style_edit(
    body: Dict[str, Any] = Body(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    summary = str(body.get("summary") or "").strip()
    if not summary:
        raise HTTPException(status_code=400, detail="Style profile text required")

    return await _save_style_profile(ctx, summary)


@router.get("/history/style")
async def rfp_history_style_get(
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    profile = await _load_style_profile(ctx)
    if not profile:
        return {"summary": "", "updated_at": None}
    return profile


@router.post("/history/style")
async def rfp_history_style_post(
    body: Optional[Dict[str, Any]] = Body(default=None),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    payload = body or {}
    model_name = str(payload.get("model") or "").strip() or settings.EKCHAT_LIGHT_TASK_MODEL

    history_excerpt = await _history_text_excerpt(ctx, max_chars=12000)
    if not history_excerpt:
        raise HTTPException(status_code=400, detail="Upload history files before building a style profile")

    system_prompt = (
        "You analyze proposal writing samples and return a style profile. "
        "Focus on tone, structure, formatting, and reusable sentence patterns."
    )
    user_prompt = f"Samples:\n{history_excerpt}\n\nReturn a concise bullet-style style profile."

    try:
        summary = await complete_chat(
            ctx.db,
            ctx.user.tenant_id,
            [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            model_name=model_name,
            max_tokens=700,
            temperature=0.2,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Style profile generation failed: {exc}") from exc

    summary = (summary or "").strip()
    if not summary:
        raise HTTPException(status_code=500, detail="Style profile generation returned empty output")

    return await _save_style_profile(ctx, summary)


@router.post("/analyze/upload")
async def rfp_analyze_upload(
    files: List[UploadFile] = File(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rag_enabled", default=False)),
):
    storage = get_storage()
    uploaded: List[Dict[str, Any]] = []

    for upload in files:
        data = await upload.read()
        if not data:
            continue

        safe_name = _safe_filename(upload.filename, default="document.bin")
        session_id = str(uuid.uuid4())
        blob_path = f"{_scope_blob_prefix(ctx, 'analyze')}/{session_id}-{safe_name}"

        await storage.upload_bytes(blob_path, data, content_type=upload.content_type)

        payload = {
            "name": safe_name,
            "mime_type": upload.content_type,
            "blob_path": blob_path,
            "file_size": len(data),
        }

        await _insert_session(
            ctx,
            kind=_ANALYZE_FILE_KIND,
            payload=payload,
            record_id=session_id,
            status_value="ready",
        )

        uploaded.append({"id": session_id, "name": safe_name, "size": len(data)})

    return {"ok": True, "uploaded": uploaded}


@router.get("/analyze/list")
async def rfp_analyze_list(
    ctx: RequestContext = Depends(require_feature("ekchat_rag_enabled", default=False)),
):
    return {"files": await _list_scope_files(ctx, "analyze")}


@router.get("/analyze/download")
async def rfp_analyze_download(
    name: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rag_enabled", default=False)),
):
    file_row = await _find_scope_file(ctx, "analyze", name=name)
    if not file_row:
        raise HTTPException(status_code=404, detail="File not found")

    storage = get_storage()
    content = await storage.download_bytes(file_row["blob_path"])
    media_type = file_row.get("mime_type") or "application/octet-stream"

    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f'inline; filename="{Path(name).name}"'},
    )


@router.get("/analyze/text")
async def rfp_analyze_text(
    name: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rag_enabled", default=False)),
):
    file_row = await _find_scope_file(ctx, "analyze", name=name)
    if not file_row:
        raise HTTPException(status_code=404, detail="File not found")

    storage = get_storage()
    content = await storage.download_bytes(file_row["blob_path"])
    text_value = _extract_text_from_bytes(name, content)

    if not text_value:
        raise HTTPException(status_code=400, detail="Could not extract readable text from the RFP file")

    return {
        "name": name,
        "text": text_value,
        "truncated": len(text_value) >= 60000,
    }


@router.post("/analyze/delete")
async def rfp_analyze_delete(
    body: Dict[str, Any] = Body(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rag_enabled", default=False)),
):
    name = str(body.get("name") or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="File name required")

    file_row = await _find_scope_file(ctx, "analyze", name=name)
    if not file_row:
        raise HTTPException(status_code=404, detail="File not found")

    storage = get_storage()
    try:
        await storage.delete_path(file_row["blob_path"])
    except Exception:
        pass

    await _delete_session(ctx, kind=_ANALYZE_FILE_KIND, session_id=file_row["id"])
    return {"ok": True, "files": await _list_scope_files(ctx, "analyze")}


@router.post("/capability-matrix/generate")
async def rfp_capability_matrix_generate(
    body: RfpCapabilityMatrixGenerateRequest,
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    rfp_name = (body.rfp_name or "").strip()
    if not rfp_name:
        raise HTTPException(status_code=400, detail="RFP file name required")

    source = await _find_scope_file(ctx, "analyze", name=rfp_name)
    if not source:
        raise HTTPException(status_code=404, detail="RFP file not found in analyze library")

    storage = get_storage()
    raw = await storage.download_bytes(source["blob_path"])
    text_value = _extract_text_from_bytes(source["name"], raw)
    if not text_value:
        raise HTTPException(status_code=400, detail="Could not extract readable text from the RFP")

    requirements = _extract_requirements(text_value, max_items=30)
    history_docs = await _load_history_evidence_docs(ctx)
    rows = _build_capability_rows(requirements, history_docs)
    rows = _normalize_capability_payload_rows(rows)

    matrix_id = str(uuid.uuid4())
    payload = {
        "rows": rows,
        "requirement_count": len(rows),
        "source": source["name"],
    }

    await ctx.db.execute(
        text(
            """
            INSERT INTO ekchat.rfp_outputs_capability_matrix (
                id,
                tenant_id,
                user_id,
                rfp_name,
                model,
                prompt_version,
                payload,
                created_at
            )
            VALUES (
                :id,
                :tenant_id,
                :user_id,
                :rfp_name,
                :model,
                :prompt_version,
                CAST(:payload AS JSONB),
                NOW()
            )
            """
        ),
        {
            "id": matrix_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "rfp_name": rfp_name,
            "model": body.model or settings.EKCHAT_LIGHT_TASK_MODEL,
            "prompt_version": body.prompt_version or "v1",
            "payload": json.dumps(payload),
        },
    )

    created_at = int(datetime.utcnow().timestamp())
    matrix_meta = {
        "id": matrix_id,
        "rfp_id": rfp_name,
        "model": body.model or settings.EKCHAT_LIGHT_TASK_MODEL,
        "prompt_version": body.prompt_version or "v1",
        "created_at": created_at,
    }

    if body.stream:
        async def stream_rows():
            yield _sse_event("init", {"rfp_id": rfp_name, "matrix": matrix_meta, "total_rows": len(rows)})
            for idx, row in enumerate(rows, start=1):
                yield _sse_event("row", {"index": idx, "total_rows": len(rows), "row": row})
            yield _sse_event("done", {"matrix": matrix_meta, "rows": rows})

        return StreamingResponse(stream_rows(), media_type="text/event-stream")

    return {"matrix": matrix_meta, "rows": rows}


@router.get("/capability-matrix/latest")
async def rfp_capability_matrix_latest(
    rfp_name: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    name = (rfp_name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="RFP file name required")

    result = await ctx.db.execute(
        text(
            """
            SELECT id, rfp_name, model, prompt_version, payload, created_at
            FROM ekchat.rfp_outputs_capability_matrix
            WHERE tenant_id = :tenant_id
              AND user_id = :user_id
              AND rfp_name = :rfp_name
            ORDER BY created_at DESC
            LIMIT 1
            """
        ),
        {
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "rfp_name": name,
        },
    )

    row = result.mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Capability matrix not found")

    payload = row.get("payload")
    if isinstance(payload, str):
        payload = json.loads(payload)
    rows = _normalize_capability_payload_rows((payload or {}).get("rows") or [])

    return {
        "matrix": {
            "id": row["id"],
            "rfp_id": row["rfp_name"],
            "model": row.get("model"),
            "prompt_version": row.get("prompt_version"),
            "created_at": _coerce_epoch_seconds(row.get("created_at")),
        },
        "rows": rows,
    }


@router.get("/capability-matrix/export")
async def rfp_capability_matrix_export(
    matrix_id: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    mid = (matrix_id or "").strip()
    if not mid:
        raise HTTPException(status_code=400, detail="Matrix id required")

    result = await ctx.db.execute(
        text(
            """
            SELECT id, rfp_name, payload
            FROM ekchat.rfp_outputs_capability_matrix
            WHERE id = :id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
            LIMIT 1
            """
        ),
        {
            "id": mid,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
        },
    )

    row = result.mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Capability matrix not found")

    payload = row.get("payload")
    if isinstance(payload, str):
        payload = json.loads(payload)
    rows = _normalize_capability_payload_rows((payload or {}).get("rows") or [])

    headers = [
        ("RFP Task / Requirement ID", "rfp_requirement_id"),
        ("Capability Area", "capability_area"),
        ("Requirement", "requirement_text"),
        ("Clause Breakdown", "clause_breakdown"),
        ("Coverage Score (0-3)", "coverage_score"),
        ("Clause-Level Findings", "rationale"),
        ("Evidence Summary", "evidence_excerpts"),
        ("Evidence Source", "evidence_sources"),
        ("Gaps / Required Fixes (exact sentences to add)", "gaps_actions"),
    ]

    data = _csv_export(rows, headers)
    filename = _safe_filename(f"{row.get('rfp_name') or 'capability-matrix'}.csv", default="capability-matrix.csv")

    return Response(
        content=data,
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/shred-document/generate")
async def rfp_shred_document_generate(
    body: RfpShredDocumentGenerateRequest,
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    rfp_name = (body.rfp_name or "").strip()
    if not rfp_name:
        raise HTTPException(status_code=400, detail="RFP file name required")

    source = await _find_scope_file(ctx, "analyze", name=rfp_name)
    if not source:
        raise HTTPException(status_code=404, detail="RFP file not found in analyze library")

    storage = get_storage()
    raw = await storage.download_bytes(source["blob_path"])
    text_value = _extract_text_from_bytes(source["name"], raw)
    if not text_value:
        raise HTTPException(status_code=400, detail="Could not extract readable text from the RFP")

    rows = _normalize_shred_payload_rows(_build_shred_rows(text_value))

    document_id = str(uuid.uuid4())
    payload = {
        "rows": rows,
        "source": source["name"],
    }

    await ctx.db.execute(
        text(
            """
            INSERT INTO ekchat.rfp_outputs_shred_document (
                id,
                tenant_id,
                user_id,
                rfp_name,
                model,
                prompt_version,
                payload,
                created_at
            )
            VALUES (
                :id,
                :tenant_id,
                :user_id,
                :rfp_name,
                :model,
                :prompt_version,
                CAST(:payload AS JSONB),
                NOW()
            )
            """
        ),
        {
            "id": document_id,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "rfp_name": rfp_name,
            "model": body.model or settings.EKCHAT_LIGHT_TASK_MODEL,
            "prompt_version": body.prompt_version or "v1",
            "payload": json.dumps(payload),
        },
    )

    return {
        "document": {
            "id": document_id,
            "rfp_id": rfp_name,
            "model": body.model or settings.EKCHAT_LIGHT_TASK_MODEL,
            "prompt_version": body.prompt_version or "v1",
            "created_at": int(datetime.utcnow().timestamp()),
        },
        "rows": rows,
    }


@router.get("/shred-document/latest")
async def rfp_shred_document_latest(
    rfp_name: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    name = (rfp_name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="RFP file name required")

    result = await ctx.db.execute(
        text(
            """
            SELECT id, rfp_name, model, prompt_version, payload, created_at
            FROM ekchat.rfp_outputs_shred_document
            WHERE tenant_id = :tenant_id
              AND user_id = :user_id
              AND rfp_name = :rfp_name
            ORDER BY created_at DESC
            LIMIT 1
            """
        ),
        {
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
            "rfp_name": name,
        },
    )

    row = result.mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Shred document not found")

    payload = row.get("payload")
    if isinstance(payload, str):
        payload = json.loads(payload)
    rows = _normalize_shred_payload_rows((payload or {}).get("rows") or [])

    return {
        "document": {
            "id": row["id"],
            "rfp_id": row["rfp_name"],
            "model": row.get("model"),
            "prompt_version": row.get("prompt_version"),
            "created_at": _coerce_epoch_seconds(row.get("created_at")),
        },
        "rows": rows,
    }


@router.get("/shred-document/export")
async def rfp_shred_document_export(
    document_id: str = Query(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    did = (document_id or "").strip()
    if not did:
        raise HTTPException(status_code=400, detail="Document id required")

    result = await ctx.db.execute(
        text(
            """
            SELECT id, rfp_name, payload
            FROM ekchat.rfp_outputs_shred_document
            WHERE id = :id
              AND tenant_id = :tenant_id
              AND user_id = :user_id
            LIMIT 1
            """
        ),
        {
            "id": did,
            "tenant_id": ctx.user.tenant_id,
            "user_id": ctx.user.user_id,
        },
    )

    row = result.mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Shred document not found")

    payload = row.get("payload")
    if isinstance(payload, str):
        payload = json.loads(payload)
    rows = _normalize_shred_payload_rows((payload or {}).get("rows") or [])

    headers = [
        ("Section", "requirement_id"),
        ("Requirement", "summary"),
    ]

    data = _csv_export(rows, headers)
    filename = _safe_filename(f"{row.get('rfp_name') or 'shred-document'}.csv", default="shred-document.csv")

    return Response(
        content=data,
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/sections/prepare")
async def rfp_sections_prepare(
    file: UploadFile = File(...),
    model: Optional[str] = Form(default=None),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    source_name = _safe_filename(file.filename, default="rfp.txt")
    text_value = _extract_text_from_bytes(source_name, raw, max_chars=_RFP_SECTION_TEXT_MAX_CHARS)
    if not text_value:
        raise HTTPException(status_code=400, detail="Could not extract readable text from uploaded file")

    profile = await _load_style_profile(ctx)
    style_summary = (profile or {}).get("summary") or ""
    history_corpus = await _history_text_excerpt(ctx, max_chars=9000)
    fallback_history_excerpt = history_corpus[:3000]
    fallback_history_words = _count_words(fallback_history_excerpt)
    model_name = (model or "").strip() or settings.EKCHAT_DEFAULT_MODEL

    prepared_text = _split_section_lines(text_value)
    outline = _extract_pws_outline(prepared_text, max_sections=_RFP_SECTION_MAX_SECTIONS)
    if not outline:
        outline = await _extract_rfp_outline(
            ctx,
            prepared_text,
            model_name=model_name,
            max_sections=min(10, _RFP_SECTION_MAX_SECTIONS),
        )
    if not outline:
        fallback_requirements = _extract_requirements(prepared_text, max_items=min(24, _RFP_SECTION_MAX_SECTIONS))
        outline = [
            {
                "title": (requirement.split(":", 1)[0][:90] if ":" in requirement else f"Section {idx}"),
                "requirements": requirement,
            }
            for idx, requirement in enumerate(fallback_requirements, start=1)
        ]
    if not outline:
        outline = [{"title": "Response", "requirements": ""}]

    sections: List[Dict[str, Any]] = []
    for idx, item in enumerate(outline, start=1):
        title = str(item.get("title") or "").strip() or f"Section {idx}"
        requirements = str(item.get("requirements") or "").strip()
        if len(title) < 6:
            title = f"Section {idx}"
        if len(title) > 140:
            title = title[:140].strip()

        history_excerpt = _select_history_excerpt(
            history_corpus,
            title,
            requirements,
            max_chars=2200,
            max_segments=3,
        )
        history_words = _count_words(history_excerpt)
        target_words = _resolve_section_target_words(history_words, fallback_history_words)
        requirement_ids = _extract_requirement_ids(requirements, idx)
        rfp_excerpt = _extract_rfp_excerpt(text_value, title, requirements, max_chars=600)
        section_history = history_excerpt or fallback_history_excerpt or "No close match found in history."

        sections.append(
            {
                "index": idx,
                "title": title,
                "requirements": requirements,
                "requirement_ids": requirement_ids,
                "target_words": target_words,
                "rfp_excerpt": rfp_excerpt,
                "history_excerpt": section_history,
                "draft": "",
                "content": "",
            }
        )

    session_payload = {
        "source_name": source_name,
        "model": model_name,
        "style_profile": style_summary,
        "sections": sections,
        "created_at": datetime.utcnow().isoformat(),
    }

    session_id = await _insert_session(
        ctx,
        kind=_SECTIONS_SESSION_KIND,
        payload=session_payload,
        expires_in_hours=24,
    )

    return {
        "session_id": session_id,
        "sections": [
            {
                "index": section["index"],
                "title": section["title"],
                "requirements": section["requirements"],
                "requirement_ids": section.get("requirement_ids") or [],
                "target_words": section["target_words"],
                "history_excerpt": section["history_excerpt"],
                "rfp_excerpt": section["rfp_excerpt"],
                "draft": section.get("draft") or "",
            }
            for section in sections
        ],
    }


@router.post("/sections/generate")
async def rfp_sections_generate(
    body: RfpSectionGenerateRequest,
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    session = await _get_sections_session(ctx, body.session_id)
    payload = session["payload"] if isinstance(session["payload"], dict) else {}
    sections = payload.get("sections") or []

    section = next((item for item in sections if int(item.get("index", -1)) == body.section_index), None)
    if not section:
        raise HTTPException(status_code=404, detail="Section not found in session")

    model_name = (body.model or payload.get("model") or settings.EKCHAT_DEFAULT_MODEL).strip()
    style_profile = payload.get("style_profile") or "No style profile available."
    requirement_ids = section.get("requirement_ids") or _extract_requirement_ids(
        str(section.get("requirements") or ""),
        int(body.section_index),
    )
    target_words_raw = section.get("target_words")
    target_words = int(target_words_raw) if isinstance(target_words_raw, int) and target_words_raw > 0 else _resolve_section_target_words(
        _count_words(str(section.get("history_excerpt") or "")),
        0,
    )
    continuity_context = _format_section_continuity(sections, body.section_index)

    system_prompt = (
        "You are an expert federal RFP response writer. "
        "Write only the requested section body and maintain continuity across sections."
    )
    user_prompt = (
        "SECTION DETAILS:\n"
        f"Title: {section.get('title')}\n"
        f"Requirements: {section.get('requirements') or 'Not specified.'}\n"
        f"Requirement IDs: {', '.join(requirement_ids) if requirement_ids else 'None'}\n"
        f"Target length: ~{target_words} words (match similar history sections; keep within about 10%).\n\n"
        "RFP EXCERPT:\n"
        f"{section.get('rfp_excerpt') or 'Not available.'}\n\n"
        "HISTORY SECTION:\n"
        "Writing style profile:\n"
        f"{style_profile or 'Not available.'}\n\n"
        "History excerpts:\n"
        f"{section.get('history_excerpt') or 'Not available.'}\n\n"
    )
    if continuity_context:
        user_prompt += (
            "PREVIOUS SECTIONS (for continuity; do not repeat):\n"
            f"{continuity_context}\n\n"
        )
    user_prompt += "Write the response for this section only. Return plain text only."

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    if body.stream:
        async def generate_stream():
            output_chunks: List[str] = []
            try:
                yield _sse_event("meta", {"session_id": body.session_id, "section_index": body.section_index})
                async for token in stream_chat_completion(
                    ctx.db,
                    ctx.user.tenant_id,
                    messages,
                    model_name=model_name,
                    max_tokens=max(900, min(2600, int(target_words * 2.2))),
                    temperature=0.2,
                ):
                    output_chunks.append(token)
                    yield _sse_event("delta", {"content": token})

                final_text = _extract_rfp_response("".join(output_chunks).strip())
                section["draft"] = final_text
                section["content"] = final_text
                section["generated_at"] = datetime.utcnow().isoformat()
                await _update_sections_payload(ctx, body.session_id, payload)
                yield _sse_event(
                    "done",
                    {
                        "content": final_text,
                        "text": final_text,
                        "coverage_ids": requirement_ids,
                        "anchor_terms_used": [],
                    },
                )
            except Exception as exc:
                yield _sse_event("error", {"message": str(exc)})

        return StreamingResponse(generate_stream(), media_type="text/event-stream")

    content = await complete_chat(
        ctx.db,
        ctx.user.tenant_id,
        messages,
        model_name=model_name,
        max_tokens=max(900, min(2600, int(target_words * 2.2))),
        temperature=0.2,
    )

    final_text = _extract_rfp_response((content or "").strip())
    if not final_text:
        raise HTTPException(status_code=500, detail="Section generation returned empty output")

    section["draft"] = final_text
    section["content"] = final_text
    section["generated_at"] = datetime.utcnow().isoformat()
    await _update_sections_payload(ctx, body.session_id, payload)

    return {
        "session_id": body.session_id,
        "text": final_text,
        "coverage_ids": requirement_ids,
        "anchor_terms_used": [],
        "section": {
            "index": section["index"],
            "title": section["title"],
            "content": final_text,
            "target_words": target_words,
        },
    }


@router.post("/response/generate")
async def rfp_response_generate(
    file: UploadFile = File(...),
    model: Optional[str] = Form(default=None),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")

    source_name = _safe_filename(file.filename, default="rfp.txt")
    text_value = _extract_text_from_bytes(source_name, raw, max_chars=18000)
    if not text_value:
        raise HTTPException(status_code=400, detail="Could not extract readable text from uploaded file")

    style_profile = (await _load_style_profile(ctx) or {}).get("summary") or "No style profile available."
    history_excerpt = await _history_text_excerpt(ctx, max_chars=3000)

    model_name = (model or "").strip() or settings.EKCHAT_DEFAULT_MODEL
    system_prompt = (
        "You are a federal proposal writer. Draft a complete response with clear sections, "
        "tailored to the provided RFP text and historical style profile."
    )
    user_prompt = (
        f"RFP source: {source_name}\n\n"
        f"RFP text:\n{text_value}\n\n"
        f"Style profile:\n{style_profile}\n\n"
        f"Historical context:\n{history_excerpt or 'Not available.'}\n\n"
        "Return a final response in markdown with section headings."
    )

    try:
        content = await complete_chat(
            ctx.db,
            ctx.user.tenant_id,
            [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            model_name=model_name,
            max_tokens=2200,
            temperature=0.2,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"RFP response generation failed: {exc}") from exc

    final_content = (content or "").strip()
    if not final_content:
        raise HTTPException(status_code=500, detail="RFP response generation returned empty output")

    response_payload = {
        "filename": source_name,
        "model": model_name,
        "content": final_content,
        "created_at": datetime.utcnow().isoformat(),
    }
    response_id = await _insert_session(
        ctx,
        kind=_RESPONSE_KIND,
        payload=response_payload,
        expires_in_hours=72,
    )

    return {
        "response_id": response_id,
        "filename": source_name,
        "content": final_content,
    }


@router.post("/response/export")
async def rfp_response_export(
    body: Dict[str, Any] = Body(...),
    ctx: RequestContext = Depends(require_feature("ekchat_rfp_enabled", default=False)),
):
    response_id = str(body.get("response_id") or "").strip()
    direct_content = str(body.get("content") or "").strip()
    fmt = str(body.get("format") or "txt").strip().lower()

    content = direct_content
    default_name = str(body.get("filename") or "rfp-response")

    if response_id:
        result = await ctx.db.execute(
            text(
                """
                SELECT payload
                FROM ekchat.rfp_sessions
                WHERE id = :id
                  AND tenant_id = :tenant_id
                  AND user_id = :user_id
                  AND kind = :kind
                LIMIT 1
                """
            ),
            {
                "id": response_id,
                "tenant_id": ctx.user.tenant_id,
                "user_id": ctx.user.user_id,
                "kind": _RESPONSE_KIND,
            },
        )

        row = result.mappings().first()
        if not row:
            raise HTTPException(status_code=404, detail="Generated response not found")

        payload = row.get("payload")
        if isinstance(payload, str):
            payload = json.loads(payload)

        payload = payload or {}
        content = str(payload.get("content") or "").strip()
        default_name = str(payload.get("filename") or default_name)

    if not content:
        raise HTTPException(status_code=400, detail="No response content available to export")

    if fmt not in {"txt", "md", "docx", "pdf"}:
        raise HTTPException(status_code=400, detail="Unsupported export format")

    file_stem = _safe_filename(default_name, default="rfp-response")
    file_stem = file_stem.rsplit(".", 1)[0]

    if fmt in {"txt", "md"}:
        media_type = "text/markdown" if fmt == "md" else "text/plain"
        filename = f"{file_stem}.{fmt}"
        return Response(
            content=content.encode("utf-8"),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if fmt == "docx":
        try:
            from docx import Document
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"DOCX export requires python-docx: {exc}") from exc

        document = Document()
        for block in content.split("\n\n"):
            document.add_paragraph(block)

        buffer = io.BytesIO()
        document.save(buffer)
        filename = f"{file_stem}.docx"
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"PDF export requires reportlab: {exc}") from exc

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    cursor_y = height - 40

    for raw_line in content.splitlines():
        line = raw_line.strip() or " "
        pdf.drawString(40, cursor_y, line[:110])
        cursor_y -= 14
        if cursor_y < 40:
            pdf.showPage()
            cursor_y = height - 40

    pdf.save()
    filename = f"{file_stem}.pdf"
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
